import csv
import glob
import os
from docx import Document
from docx.enum.table import WD_CELL_VERTICAL_ALIGNMENT
from docx.shared import Inches
from docx.shared import Pt
from robot.libraries.BuiltIn import BuiltIn


def save_screenshot_in_doc(document_file, screenshot_path, test_case_id, feature_name=None, subfeature_name=None):
    screenshot = BuiltIn().run_keyword('Capture Page Screenshot', screenshot_path)

    # Check if the document file exists
    if os.path.isfile(document_file):
        document = Document(document_file)
    else:
        # Create a new document
        document = Document()

    # Add test case ID and feature name in the same line
    test_case_line = f"Test Case ID: {test_case_id}"
    if feature_name:
        test_case_line += f"----Feature: {feature_name}"
    p = document.add_paragraph()
    run = p.add_run()
    run.bold = True
    run.underline = True
    run.add_text(test_case_line)

    if subfeature_name:
        # Add subfeature name in a new paragraph
        document.add_paragraph(f"             Subfeature: {subfeature_name}")

    screenshot_count = sum(1 for p in document.paragraphs if p.text.startswith(f"Screenshot: {feature_name}_"))

    # Generate the screenshot name with the test case ID and an incremented count
    if feature_name is not None:
        screenshot_name = f"{feature_name}_{screenshot_count + 1}"
    else:
        screenshot_name = f"{screenshot_count + 1}"

    # Add the screenshot name
    document.add_paragraph(f"Screenshot: {screenshot_name}")
    # Add the screenshot
    document.add_picture(screenshot, width=Inches(6))

    # Save the document
    document.save(document_file)

def save_screenshot_in_doc_without_testcase(document_file, screenshot_path, feature_name=None, subfeature_name=None):
    screenshot = BuiltIn().run_keyword('Capture Page Screenshot', screenshot_path)

    # Check if the document file exists
    if os.path.isfile(document_file):
        document = Document(document_file)
    else:
        # Create a new document
        document = Document()

    if subfeature_name:
        # Add subfeature name in a new paragraph
        document.add_paragraph(f"             Subfeature: {subfeature_name}")

    # Find the current screenshot count for the test case
    screenshot_count = sum(1 for p in document.paragraphs if p.text.startswith(f"Screenshot: {feature_name}_"))

    # Generate the screenshot name with the test case ID and an incremented count
    screenshot_name = f"{feature_name}_{screenshot_count + 1}"

    # Add the screenshot name
    document.add_paragraph(f"Screenshot: {screenshot_name}")
    # Add the screenshot
    document.add_picture(screenshot, width=Inches(6))

    # Save the document
    document.save(document_file)

def save_failed_screenshot_in_doc(document_file, screenshot_path, test_case_id, feature_name=None,
                                  subfeature_name=None):
    screenshot = BuiltIn().run_keyword('Capture Page Screenshot', screenshot_path)
    failed = BuiltIn().get_variable_value('${TEST_STATUS}') == 'FAIL'

    # Check if the document file exists
    if os.path.isfile(document_file):
        document = Document(document_file)
    else:
        # Create a new document
        document = Document()

    # Add test case ID and feature name in the same line
    test_case_line = f"Test Case ID: {test_case_id}"
    if feature_name:
        test_case_line += f"----Feature: {feature_name}"
    p = document.add_paragraph()
    run = p.add_run()
    run.bold = True
    run.underline = True
    run.add_text(test_case_line)

    if subfeature_name:
        # Add subfeature name in a new paragraph
        document.add_paragraph(f"             Subfeature: {subfeature_name}")

    if failed:
        # Get the error message from the test context
        error_message = BuiltIn().get_variable_value('${TEST MESSAGE}')

        # Add the error message to the document
        document.add_paragraph(f"Error Message: {error_message}")

    # Add the screenshot
    document.add_picture(screenshot, width=Inches(6))

    # Save the document
    document.save(document_file)

def get_latest_csv_file(csv_latest_file):
    csv_files = glob.glob(os.path.join(csv_latest_file, "*.csv"))
    if csv_files:
        csv_files.sort(key=os.path.getmtime, reverse=True)
    return csv_files[0]

def get_latest_jpeg_file(jpeg_latest_file):
    jpeg_files = glob.glob(os.path.join(jpeg_latest_file, "*.jpeg"))
    if jpeg_files:
        jpeg_files.sort(key=os.path.getmtime, reverse=True)
    return jpeg_files[0]

def get_latest_pdf_file(jpeg_latest_file):
    jpeg_files = glob.glob(os.path.join(jpeg_latest_file, "*.pdf"))
    if jpeg_files:
        jpeg_files.sort(key=os.path.getmtime, reverse=True)
    return jpeg_files[0]

def get_latest_file(directory):
    files = glob.glob(os.path.join(directory, "*.*"))
    if files:
        files.sort(key=os.path.getmtime, reverse=True)
    return files[0]

def add_table(document, header, table_data):
    rows = len(table_data) + 1  # Include an additional row for the header
    cols = len(header)
    table = document.add_table(rows=rows, cols=cols)
    table.style = 'Table Grid'

    # Add the header row
    header_row = table.rows[0]
    for i, column_header in enumerate(header):
        header_cell = header_row.cells[i]
        header_cell.text = column_header

        # Set the header cell's formatting
        paragraph = header_cell.paragraphs[0]
        run = paragraph.runs[0]
        run.bold = True
        run.font.size = Pt(11)  # Set the font size as desired
        header_cell.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER  # Set the vertical alignment as desired

    # Add the data rows
    for i, row_data in enumerate(table_data):
        row = table.rows[i + 1]  # Skip the header row
        for j, cell_value in enumerate(row_data):
            row.cells[j].text = str(cell_value)

def update_document_with_file_content(document_file, csv_file_path, csv_latest_file=None):
    document = Document(document_file)  # Load existing document

    if csv_latest_file:
        file_name_paragraph = document.add_paragraph()
        file_name_run = file_name_paragraph.add_run(f"File Name: {csv_latest_file}")
        file_name_run.bold = True  # Make "File Name" bold

        report_summary_paragraph = document.add_paragraph("Report Summary:")
        report_summary_paragraph.runs[0].bold = True  # Make "Report Summary" bold
        report_summary_paragraph.runs[0].font.size = Pt(11)  # Set font size of "Report Summary"

        header = ['serviceId', 'Validate Status', 'Operational Status', 'Status Reason']

        table_data = []
        with open(csv_file_path, 'r', newline='') as csvfile:
            reader = csv.DictReader(csvfile)
            for row in reader:
                row_data = [
                    row['serviceId'], row['Validate Status'], row['Operational Status'], row['Status Reason']
                ]
                if any(cell.strip() for cell in row_data):  # Check if any field in the row contains data
                    table_data.append(row_data)

        if table_data:
            add_table(document, header, table_data)
            document.add_paragraph()

    document.save(document_file)  # Save back to the same document file

def update_document_with_file_content_loyalty(document_file, csv_file_path, csv_latest_file=None):
    document = Document(document_file)  # Load existing document

    if csv_latest_file:
        file_name_paragraph = document.add_paragraph()
        file_name_run = file_name_paragraph.add_run(f"File Name: {csv_latest_file}")
        file_name_run.bold = True  # Make "File Name" bold

        report_summary_paragraph = document.add_paragraph("Report Summary:")
        report_summary_paragraph.runs[0].bold = True  # Make "Report Summary" bold
        report_summary_paragraph.runs[0].font.size = Pt(11)  # Set font size of "Report Summary"

        header = ['customerId', 'Validate Status', 'Operational Status', 'Status Reason']

        table_data = []
        with open(csv_file_path, 'r', newline='') as csvfile:
            reader = csv.DictReader(csvfile)
            for row in reader:
                row_data = [
                    row['customerId'], row['Validate Status'], row['Operational Status'], row['Status Reason']
                ]
                if any(cell.strip() for cell in row_data):  # Check if any field in the row contains data
                    table_data.append(row_data)

        if table_data:
            add_table(document, header, table_data)
            document.add_paragraph()

    document.save(document_file)  # Save back to the same document file

def update_document_with_file_content_Invoice(document_file, csv_file_path, csv_latest_file=None):
    document = Document(document_file)  # Load existing document

    if csv_latest_file:
        file_name_paragraph = document.add_paragraph()
        file_name_run = file_name_paragraph.add_run(f"File Name: {csv_latest_file}")
        file_name_run.bold = True  # Make "File Name" bold

        report_summary_paragraph = document.add_paragraph("Report Summary:")
        report_summary_paragraph.runs[0].bold = True  # Make "Report Summary" bold
        report_summary_paragraph.runs[0].font.size = Pt(11)  # Set font size of "Report Summary"

        header = ['serviceId']

        table_data = []
        with open(csv_file_path, 'r', newline='') as csvfile:
            reader = csv.DictReader(csvfile)
            for row in reader:
                row_data = [
                    row['serviceId']
                ]
                if any(cell.strip() for cell in row_data):  # Check if any field in the row contains data
                    table_data.append(row_data)

        if table_data:
            add_table(document, header, table_data)
            document.add_paragraph()

    document.save(document_file)  # Save back to the same document file

def update_document_with_file_content_jpeg(document_file, jpeg_latest_file=None):
    document = Document(document_file)  # Load existing document

    if jpeg_latest_file:
        file_name_paragraph = document.add_paragraph()
        file_name_run = file_name_paragraph.add_run(f"File Name: {jpeg_latest_file}")
        file_name_run.bold = True  # Make "File Name" bold

    document.save(document_file)

def update_document_with_file_content_pdf(document_file, pdf_latest_file=None):
    document = Document(document_file)  # Load existing document

    if pdf_latest_file:
        file_name_paragraph = document.add_paragraph()
        file_name_run = file_name_paragraph.add_run(f"File Name: {pdf_latest_file}")
        file_name_run.bold = True  # Make "File Name" bold

    document.save(document_file)

def update_document_with_file_content_failed(document_file, csv_file_path, test_case_id, csv_latest_file=None,
                                             feature_name=None, subfeature_name=None):
    if os.path.isfile(document_file):
        document = Document(document_file)
    else:
        # Create a new document
        document = Document()

    test_case_line = f"Test Case ID: {test_case_id}"
    if feature_name:
        test_case_line += f"----Feature: {feature_name}"
    p = document.add_paragraph()
    run = p.add_run()
    run.bold = True
    run.underline = True
    run.add_text(test_case_line)
    if subfeature_name:
        # Add subfeature name in a new paragraph
        document.add_paragraph(f"             Subfeature: {subfeature_name}")

    if csv_latest_file:
        file_name_paragraph = document.add_paragraph()
        file_name_run = file_name_paragraph.add_run(f"File Name: {csv_latest_file}")
        file_name_run.bold = True  # Make "File Name" bold

        report_summary_paragraph = document.add_paragraph("Report Summary:")
        report_summary_paragraph.runs[0].bold = True  # Make "Report Summary" bold
        report_summary_paragraph.runs[0].font.size = Pt(11)  # Set font size of "Report Summary"

        header = ['serviceId', 'Validate Status', 'Operational Status', 'Status Reason']

        table_data = []
        with open(csv_file_path, 'r', newline='') as csvfile:
            reader = csv.DictReader(csvfile)
            for row in reader:
                row_data = [
                    row['serviceId'], row['Validate Status'], row['Operational Status'], row['Status Reason']
                ]
                if any(cell.strip() for cell in row_data):  # Check if any field in the row contains data
                    table_data.append(row_data)

        if table_data:
            add_table(document, header, table_data)
            document.add_paragraph()

    document.save(document_file)  # Save back to the same document file

def update_document_with_file_content_failed_loyalty(document_file, csv_file_path, test_case_id, csv_latest_file=None,
                                             feature_name=None, subfeature_name=None):
    if os.path.isfile(document_file):
        document = Document(document_file)
    else:
        # Create a new document
        document = Document()

    test_case_line = f"Test Case ID: {test_case_id}"
    if feature_name:
        test_case_line += f"----Feature: {feature_name}"
    p = document.add_paragraph()
    run = p.add_run()
    run.bold = True
    run.underline = True
    run.add_text(test_case_line)
    if subfeature_name:
        # Add subfeature name in a new paragraph
        document.add_paragraph(f"             Subfeature: {subfeature_name}")

    if csv_latest_file:
        file_name_paragraph = document.add_paragraph()
        file_name_run = file_name_paragraph.add_run(f"File Name: {csv_latest_file}")
        file_name_run.bold = True  # Make "File Name" bold

        report_summary_paragraph = document.add_paragraph("Report Summary:")
        report_summary_paragraph.runs[0].bold = True  # Make "Report Summary" bold
        report_summary_paragraph.runs[0].font.size = Pt(11)  # Set font size of "Report Summary"

        header = ['customerId', 'Validate Status', 'Operational Status', 'Status Reason']

        table_data = []
        with open(csv_file_path, 'r', newline='') as csvfile:
            reader = csv.DictReader(csvfile)
            for row in reader:
                row_data = [
                    row['customerId'], row['Validate Status'], row['Operational Status'], row['Status Reason']
                ]
                if any(cell.strip() for cell in row_data):  # Check if any field in the row contains data
                    table_data.append(row_data)

        if table_data:
            add_table(document, header, table_data)
            document.add_paragraph()

    document.save(document_file)  # Save back to the same document file

def update_document_with_file_content_loyalty_batch(document_file, csv_file_path, csv_latest_file=None):
    document = Document(document_file)  # Load existing document

    if csv_latest_file:
        file_name_paragraph = document.add_paragraph()
        file_name_run = file_name_paragraph.add_run(f"File Name: {csv_latest_file}")
        file_name_run.bold = True  # Make "File Name" bold

        report_summary_paragraph = document.add_paragraph("Report Summary:")
        report_summary_paragraph.runs[0].bold = True  # Make "Report Summary" bold
        report_summary_paragraph.runs[0].font.size = Pt(11)  # Set font size of "Report Summary"

        header = ['customerId', 'Validate Status', 'Operational Status', 'Status Reason']

        table_data = []
        with open(csv_file_path, 'r', newline='') as csvfile:
            reader = csv.DictReader(csvfile)
            for row in reader:
                row_data = [
                    row['customerId'], row['Validate Status'], row['Operational Status'], row['Status Reason']
                ]
                if any(cell.strip() for cell in row_data):  # Check if any field in the row contains data
                    table_data.append(row_data)

        if table_data:
            add_table(document, header, table_data)
            document.add_paragraph()

    document.save(document_file)  # Save back to the same document file

def update_document_with_file_content_failed_loyalty_batch(document_file, csv_file_path, test_case_id,
                                                           csv_latest_file=None,
                                                           feature_name=None, subfeature_name=None):
    if os.path.isfile(document_file):
        document = Document(document_file)
    else:
        # Create a new document
        document = Document()

    test_case_line = f"Test Case ID: {test_case_id}"
    if feature_name:
        test_case_line += f"----Feature: {feature_name}"
    p = document.add_paragraph()
    run = p.add_run()
    run.bold = True
    run.underline = True
    run.add_text(test_case_line)
    if subfeature_name:
        # Add subfeature name in a new paragraph
        document.add_paragraph(f"             Subfeature: {subfeature_name}")

    if csv_latest_file:
        file_name_paragraph = document.add_paragraph()
        file_name_run = file_name_paragraph.add_run(f"File Name: {csv_latest_file}")
        file_name_run.bold = True  # Make "File Name" bold

        report_summary_paragraph = document.add_paragraph("Report Summary:")
        report_summary_paragraph.runs[0].bold = True  # Make "Report Summary" bold
        report_summary_paragraph.runs[0].font.size = Pt(11)  # Set font size of "Report Summary"

        header = ['customerId', 'Validate Status', 'Operational Status', 'Status Reason']

        table_data = []
        with open(csv_file_path, 'r', newline='') as csvfile:
            reader = csv.DictReader(csvfile)
            for row in reader:
                row_data = [
                    row['customerId'], row['Validate Status'], row['Operational Status'], row['Status Reason']
                ]
                if any(cell.strip() for cell in row_data):  # Check if any field in the row contains data
                    table_data.append(row_data)

        if table_data:
            add_table(document, header, table_data)
            document.add_paragraph()

    document.save(document_file)  # Save back to the same document file
